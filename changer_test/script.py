from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Length
from docx.oxml.ns import qn
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys

class Changer:
    def __init__(self):
        self.tarstr = None
        self.size = 10.5
        self.font = '宋体'
        self.count = 0
    def setattr(self, tarstr, size, font):
        self.tarstr = tarstr
        self.size = size
        self.font = font
    def paragraph_attribute(self, run, r = 0x00, g = 0x00, b = 0x00, bold = False):
        run.font.size = Pt(self.size)
        run.font.name = self.font
        if bold == True:
            run.font.bold = True
        run.font.color.rgb = RGBColor(r, g, b)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), self.font)

    def change_header(self, path):
        doc = Document(path)
        nol = False
        try:
            for section in doc.sections:
                section.different_first_page_header_footer=False
                header = section.header.paragraphs[0]
                #print(header.text)
                #初始有页眉
                if header.text != '':
                    #print(header.text)
                    header.clear()
                    run = header.add_run(self.tarstr)
                    self.paragraph_attribute(run)
                    header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    #print('no header')
                    #print(section.header.is_linked_to_previous)
                    section.header.is_linked_to_previous = True
                    #pass
                    #section.top_margin = Emu(0)
            #初始无页眉
            for  para in doc.paragraphs:
                oldstr = '有限公司'
                index = para.text.find(oldstr)
                if index == -1:
                    continue
                index_tail = index+4
                while (para.text[index]!=' ' or para.text[index]!=':') and index!=0:
                    index = index-1
                oldstr = para.text[index:index_tail]
                newstr = '山东汇通药业有限公司'
                #size = para.style.font.size
                #print(size)
                #left_indent = para.left_indent
                tar = para.text.replace(oldstr, newstr)
                #print(tar)
                #print(oldstr)
                #para.clear()
                #run = para.add_run(tar)
                #run.font.size = Pt(12)
                #print(para.style.font.size)
                size = Pt(12)
#                if para.style.font.size!=None:
#                    size = Pt(para.style.font.size/Length._EMUS_PER_PT)
#                print(size)
                para.text = tar
                para.style.font.size = size
                nol = True
            if nol == True:
                print('++++++++++++++++++++', path, '++++++++++++++++++')    
            doc.save(path)
            self.count = self.count+1
        except:
            print('has an error')
            pass

    def run(self, path, textBrowser=None):
        lists = os.listdir(path)
        for file in lists:
            p = os.path.join(path, file)
            if os.path.isdir(p):
                self.run(p, textBrowser)
            else:
                #print(p)
                if 'docx' in p:
                    textBrowser.append('正在修改'+p)
                    
                    self.change_header(p)
                else:
                    pass
                    

    def Usage(self):
        print('请在控制台输入python 脚本目录 修改的文件夹目录 修改后的页眉')
        print('可选参数：-Fs size 设置页眉字体大小 默认小四/10.5')
        print('可选参数：-Ff font 设置页眉字体 默认宋体')
        print('修改前请关闭所有要修改的word文档')

    
def main(argv):
    changer = Changer()
    if len(argv)<3:
        changer.Usage()
        exit()
    path = argv[1]
    tarstr = argv[2]
    size = 10.5
    font = '宋体'
    for i in range(3,len(argv),2):
        if '-Fs'== argv[i]:
            size = int(argv[i+1])
        if '-Ff'== argv[i]:
            font = argv[i+1]
    changer.setattr(tarstr, size, font)
    changer.run(path)

if __name__ == "__main__":
    main(sys.argv)
    print('修改完毕，请检查')

